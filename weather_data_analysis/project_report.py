from docx import Document
from docx.shared import Inches

# Create a Document object
doc = Document()

# Title
doc.add_heading('Project Report: Weather Data Analysis and Visualization', 0)

# Introduction
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    'The aim of this project is to analyze and visualize weather data using Python, pandas, and matplotlib. '
    'The dataset includes various weather parameters such as temperature, humidity, wind speed, precipitation, cloud cover, atmospheric pressure, UV index, season, visibility, location, and weather type. '
    'The analysis focuses on understanding the relationships between different weather parameters and visualizing these relationships through plots.'
)

# Dataset Description
doc.add_heading('Dataset Description', level=1)
doc.add_paragraph(
    'The dataset contains the following columns:\n'
    '1. Temperature: The temperature measured in degrees.\n'
    '2. Humidity: The humidity percentage.\n'
    '3. Wind Speed: The speed of the wind.\n'
    '4. Precipitation (%): The percentage of precipitation.\n'
    '5. Cloud Cover: The percentage of cloud cover.\n'
    '6. Atmospheric Pressure: The atmospheric pressure.\n'
    '7. UV Index: The UV index value.\n'
    '8. Season: The season during which the data was collected.\n'
    '9. Visibility (km): The visibility in kilometers.\n'
    '10. Location: The location of the data collection.\n'
    '11. Weather Type: The type of weather (e.g., sunny, rainy).'
)

# Data Analysis and Visualization
doc.add_heading('Data Analysis and Visualization', level=1)

# Importing Libraries and Reading Data
doc.add_heading('1. Importing Libraries and Reading Data', level=2)
doc.add_paragraph(
    '```python\n'
    'import pandas as pd\n'
    'import matplotlib.pyplot as plt\n'
    '\n'
    '# Specify the full path to your CSV file\n'
    'file_path = r"C:\\Users\\Asus.ASUS\\OneDrive\\Documents\\.vscode\\BIT_4th_sem_proffesional\\pandas\\weather_classification_data.csv"\n'
    '\n'
    '# Read the CSV file into a DataFrame\n'
    'weather_df = pd.read_csv(file_path)\n'
    '```'
)

# Displaying Basic Statistics
doc.add_heading('2. Displaying Basic Statistics', level=2)
doc.add_paragraph(
    'Basic statistical details provide insights into the dataset:\n'
    '```python\n'
    'print("Basic Statistics of the DataFrame:")\n'
    'print(weather_df.describe())\n'
    '```'
)

# Printing Column Names
doc.add_heading('3. Printing Column Names', level=2)
doc.add_paragraph(
    'To verify the column names in the dataset:\n'
    '```python\n'
    'print("\\nColumn Names:")\n'
    'print(weather_df.columns)\n'
    '```'
)

# Visualizing Temperature by Season
doc.add_heading('4. Visualizing Temperature by Season', level=2)
doc.add_paragraph(
    'This plot shows the variation of temperature across different seasons.\n'
    '```python\n'
    'plt.figure(figsize=(10, 6))\n'
    'plt.plot(weather_df["Season"], weather_df["Temperature"], marker="o", linestyle="-", label="Temperature")\n'
    'plt.xlabel("Season")\n'
    'plt.ylabel("Temperature")\n'
    'plt.title("Temperature by Season")\n'
    'plt.legend()\n'
    'plt.grid(True)\n'
    'plt.xticks(rotation=45)\n'
    'plt.tight_layout()\n'
    'plt.savefig("temperature_by_season.png")\n'
    'plt.show()\n'
    '```'
)

# Visualizing Temperature vs. Humidity
doc.add_heading('5. Visualizing Temperature vs. Humidity', level=2)
doc.add_paragraph(
    'This scatter plot depicts the relationship between temperature and humidity.\n'
    '```python\n'
    'plt.figure(figsize=(10, 6))\n'
    'plt.scatter(weather_df["Humidity"], weather_df["Temperature"], label="Temperature vs Humidity", color="r")\n'
    'plt.xlabel("Humidity")\n'
    'plt.ylabel("Temperature")\n'
    'plt.title("Temperature vs. Humidity")\n'
    'plt.legend()\n'
    'plt.grid(True)\n'
    'plt.tight_layout()\n'
    'plt.savefig("temperature_vs_humidity.png")\n'
    'plt.show()\n'
    '```'
)

# Visualizing Temperature vs. Wind Speed
doc.add_heading('6. Visualizing Temperature vs. Wind Speed', level=2)
doc.add_paragraph(
    'This scatter plot shows how temperature varies with wind speed.\n'
    '```python\n'
    'plt.figure(figsize=(10, 6))\n'
    'plt.scatter(weather_df["Wind Speed"], weather_df["Temperature"], label="Temperature vs Wind Speed", color="g")\n'
    'plt.xlabel("Wind Speed")\n'
    'plt.ylabel("Temperature")\n'
    'plt.title("Temperature vs. Wind Speed")\n'
    'plt.legend()\n'
    'plt.grid(True)\n'
    'plt.tight_layout()\n'
    'plt.savefig("temperature_vs_wind_speed.png")\n'
    'plt.show()\n'
    '```'
)

# Visualizing Temperature vs. UV Index
doc.add_heading('7. Visualizing Temperature vs. UV Index', level=2)
doc.add_paragraph(
    'This scatter plot illustrates the relationship between temperature and the UV index.\n'
    '```python\n'
    'plt.figure(figsize=(10, 6))\n'
    'plt.scatter(weather_df["UV Index"], weather_df["Temperature"], label="Temperature vs UV Index", color="b")\n'
    'plt.xlabel("UV Index")\n'
    'plt.ylabel("Temperature")\n'
    'plt.title("Temperature vs. UV Index")\n'
    'plt.legend()\n'
    'plt.grid(True)\n'
    'plt.tight_layout()\n'
    'plt.savefig("temperature_vs_uv_index.png")\n'
    'plt.show()\n'
    '```'
)

# Sum of Missing Values in Each Column
doc.add_heading('8. Sum of Missing Values in Each Column', level=2)
doc.add_paragraph(
    'Checking for missing values helps in understanding data completeness.\n'
    '```python\n'
    'print("\\nSum of Missing Values in Each Column:")\n'
    'print(weather_df.isnull().sum())\n'
    '```'
)

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'This project provided a comprehensive analysis and visualization of weather data. '
    'The visualizations helped in understanding how temperature varies with different weather parameters such as season, humidity, wind speed, and UV index. '
    'The analysis can be extended to explore more relationships and perform deeper data analysis.\n\n'
    'The project also demonstrated the use of pandas for data manipulation and matplotlib for data visualization, highlighting the power of Python in handling and visualizing complex datasets.'
)

# Future Work
doc.add_heading('Future Work', level=1)
doc.add_paragraph(
    'Future work could involve:\n'
    '- Exploring more complex relationships between different weather parameters.\n'
    '- Implementing machine learning models to predict weather conditions based on the available data.\n'
    '- Enhancing visualizations with interactive plots using libraries such as Plotly or Bokeh.'
)

# Save the document
file_path = r"C:\Users\Asus.ASUS\OneDrive\Documents\.vscode\BIT_4th_sem_proffesional\Weather_Data_Analysis_Project_Report.docx"
doc.save(file_path)
doc.save(file_path)

file_path
